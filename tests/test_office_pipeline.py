import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import io
import pytest


def _make_word_bytes_simple() -> bytes:
    """Build a minimal .docx with headings and paragraphs."""
    from docx import Document
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Hello world paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Content in section two.")
    doc.core_properties.title = "Test Doc"
    doc.core_properties.author = "Test Author"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_word_returns_sections():
    from office_pipeline import extract_text_from_word
    sections, metadata = extract_text_from_word(_make_word_bytes_simple())
    assert len(sections) >= 1
    all_text = " ".join(text for _, text in sections)
    assert "Hello world paragraph" in all_text
    assert "Content in section two" in all_text


def test_extract_word_section_numbers_are_ints():
    from office_pipeline import extract_text_from_word
    sections, _ = extract_text_from_word(_make_word_bytes_simple())
    for num, _ in sections:
        assert isinstance(num, int)


def test_extract_word_metadata_has_required_keys():
    from office_pipeline import extract_text_from_word
    _, metadata = extract_text_from_word(_make_word_bytes_simple())
    for key in ("comments", "tracked_changes", "footnotes", "endnotes", "properties"):
        assert key in metadata, f"Missing key: {key}"


def test_extract_word_properties():
    from office_pipeline import extract_text_from_word
    _, metadata = extract_text_from_word(_make_word_bytes_simple())
    props = metadata["properties"]
    assert props["title"] == "Test Doc"
    assert props["author"] == "Test Author"
    for key in ("created", "modified", "subject", "keywords", "description"):
        assert key in props


def test_extract_word_empty_annotations_are_lists():
    from office_pipeline import extract_text_from_word
    _, metadata = extract_text_from_word(_make_word_bytes_simple())
    assert metadata["comments"] == []
    assert metadata["tracked_changes"] == []
    assert metadata["footnotes"] == []
    assert metadata["endnotes"] == []


def _make_word_bytes_with_comments() -> bytes:
    """Build a .docx that contains a comment using lxml directly."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This text has a comment.")

    # Build comments XML part manually
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:comment w:id="1" w:author="Alice" w:date="2026-01-01T00:00:00Z">'
        '<w:p><w:r><w:t>Great point!</w:t></w:r></w:p>'
        '</w:comment>'
        '</w:comments>'
    )

    # Inject comments part into the docx package
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    comments_part = Part(
        PackURI("/word/comments.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        comments_xml.encode(),
        doc.part.package,
    )
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    doc.part.relate_to(comments_part, rel_type)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_word_bytes_with_tracked_changes() -> bytes:
    """Build a .docx with a tracked insertion."""
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    para = doc.add_paragraph()

    # Add a normal run
    run = para.add_run("Original text. ")

    # Add a tracked insertion run
    ins_xml = (
        '<w:ins xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:id="1" w:author="Bob" w:date="2026-01-02T00:00:00Z">'
        '<w:r><w:t>Inserted text.</w:t></w:r>'
        '</w:ins>'
    )
    ins_el = etree.fromstring(ins_xml)
    para._element.append(ins_el)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_word_comments():
    from office_pipeline import extract_text_from_word
    _, metadata = extract_text_from_word(_make_word_bytes_with_comments())
    assert len(metadata["comments"]) == 1
    c = metadata["comments"][0]
    assert c["author"] == "Alice"
    assert "Great point" in c["text"]
    assert "date" in c
    assert "ref_text" in c


def test_extract_word_tracked_changes():
    from office_pipeline import extract_text_from_word
    _, metadata = extract_text_from_word(_make_word_bytes_with_tracked_changes())
    inserts = [tc for tc in metadata["tracked_changes"] if tc["type"] == "insert"]
    assert len(inserts) >= 1
    assert inserts[0]["author"] == "Bob"
    assert "Inserted text" in inserts[0]["text"]


def _make_excel_bytes_simple() -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Results"
    ws["A1"] = "Name"
    ws["B1"] = "Score"
    ws["A2"] = "Alice"
    ws["B2"] = 95
    ws["A3"] = "Bob"
    ws["B3"] = 87
    ws2 = wb.create_sheet("Summary")
    ws2["A1"] = "Total"
    ws2["B1"] = 2
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_extract_excel_returns_sheets():
    from office_pipeline import extract_text_from_excel
    sheets, metadata = extract_text_from_excel(_make_excel_bytes_simple())
    assert len(sheets) == 2
    names = [name for name, _ in sheets]
    assert "Results" in names
    assert "Summary" in names


def test_extract_excel_cell_values_in_text():
    from office_pipeline import extract_text_from_excel
    sheets, _ = extract_text_from_excel(_make_excel_bytes_simple())
    results_text = next(text for name, text in sheets if name == "Results")
    assert "Alice" in results_text
    assert "95" in results_text
    assert "Bob" in results_text


def test_extract_excel_metadata_has_required_keys():
    from office_pipeline import extract_text_from_excel
    _, metadata = extract_text_from_excel(_make_excel_bytes_simple())
    for key in ("cell_notes", "threaded_comments", "formulas", "embedded_objects", "properties"):
        assert key in metadata, f"Missing key: {key}"


def test_extract_excel_empty_annotations_are_lists():
    from office_pipeline import extract_text_from_excel
    _, metadata = extract_text_from_excel(_make_excel_bytes_simple())
    assert metadata["cell_notes"] == []
    assert metadata["threaded_comments"] == []
    assert metadata["embedded_objects"] == []


def _make_excel_bytes_with_formulas() -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Calc"
    ws["A1"] = 10
    ws["A2"] = 20
    ws["A3"] = "=A1+A2"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_excel_bytes_with_cell_note() -> bytes:
    from openpyxl import Workbook
    from openpyxl.comments import Comment
    wb = Workbook()
    ws = wb.active
    ws.title = "Notes"
    ws["A1"] = "Value"
    comment = Comment("Remember to update this quarterly.", "Finance Team")
    ws["A1"].comment = comment
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_extract_excel_formulas():
    from office_pipeline import extract_text_from_excel
    _, metadata = extract_text_from_excel(_make_excel_bytes_with_formulas())
    formulas = metadata["formulas"]
    assert len(formulas) >= 1
    f = next(x for x in formulas if x["formula"] == "=A1+A2")
    assert f["sheet"] == "Calc"
    assert f["cell"] == "A3"
    assert "formula" in f
    assert "cached_value" in f


def test_extract_excel_cell_notes():
    from office_pipeline import extract_text_from_excel
    _, metadata = extract_text_from_excel(_make_excel_bytes_with_cell_note())
    notes = metadata["cell_notes"]
    assert len(notes) >= 1
    n = notes[0]
    assert n["sheet"] == "Notes"
    assert "A1" in n["cell"]
    assert "Finance Team" in n["author"]
    assert "quarterly" in n["text"]


def test_ingest_office_word_shape():
    from office_pipeline import ingest_office
    doc = ingest_office(_make_word_bytes_simple(), "test.docx")
    assert doc is not None
    assert "doc_id" in doc
    assert "filename" in doc
    assert "page_count" in doc
    assert "chunks" in doc
    assert "embeddings" in doc
    assert doc["filename"] == "test.docx"
    assert len(doc["chunks"]) >= 1
    assert doc["embeddings"].shape[0] == len(doc["chunks"])


def test_ingest_office_excel_shape():
    from office_pipeline import ingest_office
    doc = ingest_office(_make_excel_bytes_simple(), "data.xlsx")
    assert doc is not None
    assert doc["filename"] == "data.xlsx"
    assert len(doc["chunks"]) >= 1
    assert doc["embeddings"].shape[0] == len(doc["chunks"])


def test_ingest_office_unknown_extension_returns_none():
    from office_pipeline import ingest_office
    result = ingest_office(b"garbage", "file.txt")
    assert result is None


def test_ingest_office_empty_word_returns_none():
    from office_pipeline import ingest_office
    from docx import Document
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    result = ingest_office(buf.getvalue(), "empty.docx")
    assert result is None


@pytest.mark.asyncio
async def test_read_word_tool_returns_text(tmp_path):
    from agent_utils import _read_word
    import os
    # Write inside project sandbox (cwd)
    project_tmp = os.path.join(os.getcwd(), "tmp_test_word.docx")
    try:
        with open(project_tmp, "wb") as f:
            f.write(_make_word_bytes_simple())
        result = await _read_word(project_tmp)
        assert "ERROR" not in result
        assert "Hello world paragraph" in result
    finally:
        if os.path.exists(project_tmp):
            os.remove(project_tmp)


@pytest.mark.asyncio
async def test_read_excel_tool_returns_text(tmp_path):
    from agent_utils import _read_excel
    import os
    # Write inside project sandbox (cwd)
    project_tmp = os.path.join(os.getcwd(), "tmp_test_excel.xlsx")
    try:
        with open(project_tmp, "wb") as f:
            f.write(_make_excel_bytes_simple())
        result = await _read_excel(project_tmp)
        assert "ERROR" not in result
        assert "Alice" in result
    finally:
        if os.path.exists(project_tmp):
            os.remove(project_tmp)


@pytest.mark.asyncio
async def test_read_word_tool_missing_file():
    from agent_utils import _read_word
    result = await _read_word("/tmp/does_not_exist_abc123.docx")
    assert result.startswith("ERROR")


@pytest.mark.asyncio
async def test_read_excel_tool_missing_file():
    from agent_utils import _read_excel
    result = await _read_excel("/tmp/does_not_exist_abc123.xlsx")
    assert result.startswith("ERROR")
