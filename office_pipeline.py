import io
import uuid
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn


def extract_text_from_word(file_bytes: bytes) -> tuple[list[tuple[int, str]], dict]:
    """Extract body text and all annotations from a .docx file.

    Returns:
        (sections, metadata) where sections is list of (section_num, text)
        and metadata contains comments, tracked_changes, footnotes, endnotes, properties.
    """
    doc = Document(io.BytesIO(file_bytes))

    # --- Body text, split at heading boundaries ---
    sections: list[tuple[int, str]] = []
    section_num = 0
    current_lines: list[str] = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_lines:
                section_num += 1
                sections.append((section_num, "\n".join(current_lines)))
                current_lines = []
            section_num += 1
            if para.text.strip():
                current_lines.append(para.text)
        else:
            if para.text.strip():
                current_lines.append(para.text)

    if current_lines:
        section_num += 1
        sections.append((section_num, "\n".join(current_lines)))

    # If no headings produced sections, return body as single section
    if not sections:
        body_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if body_text:
            sections = [(1, body_text)]

    # --- Comments ---
    comments = _extract_word_comments(doc)

    # --- Tracked changes ---
    tracked_changes = _extract_tracked_changes(doc)

    # --- Footnotes ---
    footnotes = _extract_word_notes(doc, "footnotes")

    # --- Endnotes ---
    endnotes = _extract_word_notes(doc, "endnotes")

    # --- Document properties ---
    props = doc.core_properties
    # Note: python-docx exposes dc:description as .comments in some versions,
    # so we check .description first, then fall back to .comments if needed.
    properties = {
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "description": getattr(props, "description", None) or getattr(props, "comments", None) or "",
    }

    metadata = {
        "comments": comments,
        "tracked_changes": tracked_changes,
        "footnotes": footnotes,
        "endnotes": endnotes,
        "properties": properties,
    }
    return sections, metadata


def _extract_word_comments(doc) -> list[dict]:
    comments = []
    _COMMENTS_REL = (
        "http://schemas.openxmlformats.org/officeDocument/2006/"
        "relationships/comments"
    )
    try:
        part = doc.part.part_related_by(_COMMENTS_REL)
        for c in part._element.findall(qn("w:comment")):
            author = c.get(qn("w:author"), "")
            date = c.get(qn("w:date"), "")
            text = "".join(t.text for t in c.iter(qn("w:t")) if t.text)
            comments.append({"author": author, "date": date, "text": text, "ref_text": ""})
    except KeyError:
        pass
    return comments


def _extract_tracked_changes(doc) -> list[dict]:
    changes = []
    body = doc.element.body
    for ins in body.iter(qn("w:ins")):
        author = ins.get(qn("w:author"), "")
        date = ins.get(qn("w:date"), "")
        text = "".join(t.text for t in ins.iter(qn("w:t")) if t.text)
        if text:
            changes.append({"author": author, "date": date, "type": "insert", "text": text})
    for del_ in body.iter(qn("w:del")):
        author = del_.get(qn("w:author"), "")
        date = del_.get(qn("w:date"), "")
        text = "".join(t.text for t in del_.iter(qn("w:delText")) if t.text)
        if text:
            changes.append({"author": author, "date": date, "type": "delete", "text": text})
    return changes


def _extract_word_notes(doc, note_type: str) -> list[dict]:
    """note_type is 'footnotes' or 'endnotes'."""
    notes = []
    rel = (
        f"http://schemas.openxmlformats.org/officeDocument/2006/"
        f"relationships/{note_type}"
    )
    tag = qn("w:footnote") if note_type == "footnotes" else qn("w:endnote")
    skip_types = {"separator", "continuationSeparator", "continuationNotice"}
    try:
        part = doc.part.part_related_by(rel)
        for el in part._element.findall(tag):
            if el.get(qn("w:type"), "") in skip_types:
                continue
            note_id = el.get(qn("w:id"), "")
            text = "".join(t.text for t in el.iter(qn("w:t")) if t.text)
            if text:
                notes.append({"ref_num": note_id, "text": text})
    except KeyError:
        pass
    return notes


def extract_text_from_excel(file_bytes: bytes) -> tuple[list[tuple[str, str]], dict]:
    """Extract cell text and all annotations from a .xlsx file.

    Returns:
        (sheets, metadata) where sheets is list of (sheet_name, tsv_text)
        and metadata contains cell_notes, threaded_comments, formulas,
        embedded_objects, properties.
    """
    import openpyxl

    # Load twice: once for formulas, once for cached values
    wb_formula = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=False)
    wb_cached = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)

    sheets: list[tuple[str, str]] = []
    cell_notes: list[dict] = []
    formulas: list[dict] = []

    for sheet_name in wb_cached.sheetnames:
        ws_cached = wb_cached[sheet_name]
        ws_formula = wb_formula[sheet_name]

        rows = []
        for row in ws_cached.iter_rows():
            parts = []
            for cell in row:
                if cell.value is not None:
                    parts.append(f"{cell.coordinate}={cell.value}")
            if parts:
                rows.append("\t".join(parts))

        if rows:
            sheets.append((sheet_name, "\n".join(rows)))

        # Capture formulas (cells whose formula-mode value starts with "=")
        for row in ws_formula.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    cached_cell = ws_cached[cell.coordinate]
                    formulas.append({
                        "sheet": sheet_name,
                        "cell": cell.coordinate,
                        "formula": cell.value,
                        "cached_value": str(cached_cell.value) if cached_cell.value is not None else "",
                    })

        # Legacy cell notes (comments) — iterate cells to find per-cell comments
        for row in ws_cached.iter_rows():
            for cell in row:
                comment = cell.comment
                if comment is not None:
                    cell_notes.append({
                        "sheet": sheet_name,
                        "cell": cell.coordinate,
                        "author": comment.author or "",
                        "text": str(comment.text) if comment.text else "",
                    })

    threaded_comments = _extract_excel_threaded_comments(file_bytes)
    embedded_objects = _extract_excel_embedded_objects(file_bytes)

    # VBA macro detection
    if wb_formula.vba_archive is not None:
        embedded_objects.append({"type": "vba_macro", "flagged": True, "extracted_text": ""})

    props = wb_cached.properties
    properties = {
        "title": props.title or "" if props else "",
        "author": props.creator or "" if props else "",
        "created": str(props.created) if props and props.created else "",
        "modified": str(props.modified) if props and props.modified else "",
        "keywords": props.keywords or "" if props else "",
    }

    metadata = {
        "cell_notes": cell_notes,
        "threaded_comments": threaded_comments,
        "formulas": formulas,
        "embedded_objects": embedded_objects,
        "properties": properties,
    }
    return sheets, metadata


def _extract_excel_threaded_comments(file_bytes: bytes) -> list[dict]:
    results = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            thread_files = [n for n in z.namelist() if "threadedComment" in n and n.endswith(".xml")]
            for tf in thread_files:
                ns = {"tc": "http://schemas.microsoft.com/office/spreadsheetml/2018/threadedcomments"}
                root = ET.parse(z.open(tf)).getroot()

                # First pass: collect all entries by their own id
                all_entries: dict[str, dict] = {}
                for tc in root.findall(".//tc:threadedComment", ns):
                    tc_id = tc.get("id", "")
                    ref = tc.get("ref", "")
                    parent_id = tc.get("parentId", "")
                    author_id = tc.get("personId", "")
                    text_el = tc.find("tc:text", ns)
                    text = text_el.text if text_el is not None else ""
                    all_entries[tc_id] = {
                        "id": tc_id,
                        "ref": ref,
                        "parent_id": parent_id,
                        "entry": {"author_id": author_id, "text": text or ""},
                    }

                # Second pass: build thread structure
                threads: dict[str, dict] = {}
                for tc_id, item in all_entries.items():
                    parent_id = item["parent_id"]
                    if not parent_id:
                        threads[tc_id] = {"cell": item["ref"], "thread": [item["entry"]]}
                for tc_id, item in all_entries.items():
                    parent_id = item["parent_id"]
                    if parent_id and parent_id in threads:
                        threads[parent_id]["thread"].append(item["entry"])

                results.extend(threads.values())
    except Exception:
        pass
    return results


def _extract_excel_embedded_objects(file_bytes: bytes) -> list[dict]:
    results = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            embedding_files = [n for n in z.namelist() if "embeddings/" in n]
            for ef in embedding_files:
                ext = ef.rsplit(".", 1)[-1].lower() if "." in ef else ""
                data = z.read(ef)
                if ext in ("docx", "xlsx", "pptx"):
                    results.append({"type": f"embedded_{ext}", "extracted_text": f"[embedded {ext} — binary]"})
                elif ext in ("bin", "ole", ""):
                    try:
                        import olefile
                        if olefile.isOleFile(io.BytesIO(data)):
                            ole = olefile.OleFileIO(io.BytesIO(data))
                            if ole.exists("WordDocument"):
                                obj_type = "embedded_word_ole"
                            elif ole.exists("Workbook"):
                                obj_type = "embedded_excel_ole"
                            else:
                                obj_type = "embedded_ole"
                            results.append({"type": obj_type, "extracted_text": f"[OLE object: {ef}]"})
                    except Exception:
                        results.append({"type": "embedded_binary", "extracted_text": f"[binary object: {ef}]"})
    except Exception:
        pass
    return results


def ingest_office(file_bytes: bytes, filename: str) -> dict | None:
    """Ingest a Word or Excel file into the RAG document store.

    Returns the same dict shape as pdf_pipeline.ingest_pdf, or None if
    the file type is unrecognised or produces no text.
    """
    from pdf_pipeline import chunk_text, embed_texts

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "docx":
        sections, metadata = extract_text_from_word(file_bytes)
        pages = list(sections)
        # Only include metadata page when there is body content
        if pages:
            extra = _format_word_metadata_text(metadata)
            if extra:
                pages.append((len(pages) + 1, extra))
    elif ext == "xlsx":
        sheets, metadata = extract_text_from_excel(file_bytes)
        pages = [(i + 1, text) for i, (_, text) in enumerate(sheets)]
        extra = _format_excel_metadata_text(metadata)
        if extra:
            pages.append((len(pages) + 1, extra))
    else:
        return None

    if not pages:
        return None

    chunks = chunk_text(pages)
    if not chunks:
        return None

    texts = [c["text"] for c in chunks]
    embeddings, emb_latency_ms = embed_texts(texts)

    doc_id = uuid.uuid4().hex[:8]
    page_count = pages[-1][0] if pages else 0
    return {
        "doc_id": doc_id,
        "filename": filename,
        "page_count": page_count,
        "chunks": chunks,
        "embeddings": embeddings,
        "embedding_latency_ms": emb_latency_ms
    }


def _format_word_metadata_text(metadata: dict) -> str:
    parts = []
    if metadata.get("comments"):
        parts.append("## Comments")
        for c in metadata["comments"]:
            parts.append(f"[{c['author']}, {c['date']}]: {c['text']}")
    if metadata.get("tracked_changes"):
        parts.append("## Tracked Changes")
        for tc in metadata["tracked_changes"]:
            parts.append(f"[{tc['type'].upper()} by {tc['author']} on {tc['date']}]: {tc['text']}")
    if metadata.get("footnotes"):
        parts.append("## Footnotes")
        for fn in metadata["footnotes"]:
            parts.append(f"[{fn['ref_num']}]: {fn['text']}")
    if metadata.get("endnotes"):
        parts.append("## Endnotes")
        for en in metadata["endnotes"]:
            parts.append(f"[{en['ref_num']}]: {en['text']}")
    if metadata.get("properties"):
        p = metadata["properties"]
        kv = ", ".join(f"{k}: {v}" for k, v in p.items() if v)
        if kv:
            parts.append(f"## Document Properties\n{kv}")
    return "\n".join(parts)


def _format_excel_metadata_text(metadata: dict) -> str:
    parts = []
    if metadata.get("cell_notes"):
        parts.append("## Cell Notes")
        for n in metadata["cell_notes"]:
            parts.append(f"[{n['sheet']}!{n['cell']}, {n['author']}]: {n['text']}")
    if metadata.get("threaded_comments"):
        parts.append("## Threaded Comments")
        for tc in metadata["threaded_comments"]:
            for reply in tc.get("thread", []):
                parts.append(f"[{tc['cell']}]: {reply['text']}")
    if metadata.get("formulas"):
        parts.append("## Formulas")
        for f in metadata["formulas"]:
            parts.append(f"[{f['sheet']}!{f['cell']}] {f['formula']} = {f['cached_value']}")
    if metadata.get("embedded_objects"):
        parts.append("## Embedded Objects")
        for obj in metadata["embedded_objects"]:
            if obj.get("flagged"):
                parts.append("WARNING: VBA macro detected (not executed)")
            else:
                parts.append(f"[{obj['type']}]: {obj['extracted_text']}")
    if metadata.get("properties"):
        p = metadata["properties"]
        kv = ", ".join(f"{k}: {v}" for k, v in p.items() if v)
        if kv:
            parts.append(f"## Document Properties\n{kv}")
    return "\n".join(parts)


def write_word_document(path, spec: dict) -> None:
    """Create a .docx file from a structured spec dict."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    props = spec.get("properties", {})
    core = doc.core_properties
    if props.get("title"):
        core.title = props["title"]
    if props.get("author"):
        core.author = props["author"]
    if props.get("subject"):
        core.subject = props["subject"]
    if props.get("keywords"):
        core.keywords = props["keywords"]

    for section in spec.get("sections", []):
        stype = section.get("type", "paragraph")

        if stype == "heading":
            level = min(max(int(section.get("level", 1)), 1), 6)
            doc.add_heading(section.get("text", ""), level=level)

        elif stype == "paragraph":
            p = doc.add_paragraph()
            run = p.add_run(section.get("text", ""))
            if section.get("bold"):
                run.bold = True
            if section.get("italic"):
                run.italic = True
            if section.get("underline"):
                run.underline = True

        elif stype == "table":
            rows_data = section.get("rows", [])
            if rows_data:
                n_cols = max((len(r) for r in rows_data), default=1)
                table = doc.add_table(rows=len(rows_data), cols=n_cols)
                for i, row in enumerate(rows_data):
                    for j, cell_val in enumerate(row):
                        if j < n_cols:
                            table.cell(i, j).text = str(cell_val) if cell_val is not None else ""
                for merge in section.get("merge", []):
                    r, c = merge["row"], merge["col"]
                    rs, cs = merge["rowspan"], merge["colspan"]
                    table.cell(r, c).merge(table.cell(r + rs - 1, c + cs - 1))

        elif stype in ("footnote", "endnote"):
            # python-docx has no public footnote/endnote API; render as labelled paragraph
            label = "Footnote" if stype == "footnote" else "Endnote"
            p = doc.add_paragraph()
            run = p.add_run(f"[{label} {section.get('ref_paragraph', '')}]: {section.get('text', '')}")
            run.font.size = Pt(9)

    doc.save(str(path))


def write_excel_document(path, spec: dict) -> None:
    """Create a .xlsx file from a structured spec dict."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.utils.cell import range_boundaries

    wb = Workbook()
    wb.remove(wb.active)

    props = spec.get("properties", {})
    if props.get("title"):
        wb.properties.title = props["title"]
    if props.get("author"):
        wb.properties.creator = props["author"]
    if props.get("keywords"):
        wb.properties.keywords = props["keywords"]

    for sheet_spec in spec.get("sheets", []):
        ws = wb.create_sheet(title=sheet_spec.get("name", "Sheet"))
        rows_data = sheet_spec.get("rows", [])

        for r_idx, row in enumerate(rows_data, start=1):
            for c_idx, cell_spec in enumerate(row, start=1):
                if cell_spec is None:
                    continue
                cell = ws.cell(row=r_idx, column=c_idx)
                if isinstance(cell_spec, dict):
                    val = cell_spec.get("formula") if "formula" in cell_spec else cell_spec.get("value")
                    cell.value = val
                    font_kw: dict = {}
                    if cell_spec.get("bold"):
                        font_kw["bold"] = True
                    if cell_spec.get("italic"):
                        font_kw["italic"] = True
                    if font_kw:
                        cell.font = Font(**font_kw)
                    if cell_spec.get("border"):
                        from openpyxl.styles import Border, Side
                        _side = Side(border_style="thin")
                        cell.border = Border(left=_side, right=_side, top=_side, bottom=_side)
                    if cell_spec.get("bg_color"):
                        cell.fill = PatternFill(fill_type="solid", fgColor=cell_spec["bg_color"])
                    if cell_spec.get("number_format"):
                        cell.number_format = cell_spec["number_format"]
                    if cell_spec.get("alignment"):
                        cell.alignment = Alignment(horizontal=cell_spec["alignment"])
                else:
                    cell.value = cell_spec

        for merge_range in sheet_spec.get("merges", []):
            ws.merge_cells(merge_range)

        for chart_spec in sheet_spec.get("charts", []):
            ctype = chart_spec.get("type", "bar")
            if ctype == "bar":
                chart = BarChart()
            elif ctype == "line":
                chart = LineChart()
            elif ctype == "pie":
                chart = PieChart()
            else:
                continue
            chart.title = chart_spec.get("title", "")
            data_range = chart_spec.get("data_range", "")
            if data_range:
                cell_range = data_range.split("!")[-1]
                min_col, min_row, max_col, max_row = range_boundaries(cell_range)
                data = Reference(ws, min_col=min_col, min_row=min_row, max_col=max_col, max_row=max_row)
                chart.add_data(data, titles_from_data=True)
            ws.add_chart(chart, chart_spec.get("anchor", "E1"))

    if not wb.worksheets:
        wb.create_sheet("Sheet1")

    wb.save(str(path))


def format_office_read_output(sections_or_sheets: list[tuple], metadata: dict, filetype: str) -> str:
    """Format extraction results into a single readable string for the agent."""
    parts = []
    label = "Section" if filetype == "word" else "Sheet"
    for key, text in sections_or_sheets:
        parts.append(f"[{label} {key}]\n{text}")
    annotation_text = (
        _format_word_metadata_text(metadata)
        if filetype == "word"
        else _format_excel_metadata_text(metadata)
    )
    if annotation_text:
        parts.append(annotation_text)
    result = "\n\n".join(parts)
    return result[:15000] if len(result) > 15000 else result
